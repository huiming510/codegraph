# -*- coding: utf-8 -*-
"""
# @Time    : 2025/11/7 16:54
# @Author  : cuils
# @Description:
"""

import uuid
import base64
from typing import List
from pathlib import Path
from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

from .base_parser import BaseParser
from ..base import LinkDocument
from ..format_convert import BaseFormatConverter, LibreofficeFormatConverter


class DocxParser:
    """解析docx文件，并转换为markdown"""
    def __init__(self, docx_path):
        docx_path = Path(docx_path)
        assert docx_path.suffix == ".docx", "文件必须为docx格式，请先转为docx格式。"
        self.docx_path = docx_path
        self.doc = Document(self.docx_path)
        self.image_counter = 0
        self.doc_images = []
        self.elements = [] # 存储所有元素及位置信息
        self.position = 0 # 当前索引位置

    def get_document_properties(self):
        """获取文档属性"""
        core_props = self.doc.core_properties
        return {
            "title": core_props.title, # 标题
            "author": core_props.author, # 作者
            "created": core_props.created, # 创建时间
            "modified": core_props.modified, # 修改时间
            "subject": core_props.subject, # 主题
            "keywords": core_props.keywords # 关键词
        }

    def parse_all(self):
        """解析文档"""
        for _element in self.doc.element.body:
            if isinstance(_element, CT_P):
                paragraph = Paragraph(_element, self.doc)
                p_element = self.parse_paragraph(paragraph)
                self.elements.append(p_element)

            elif isinstance(_element, CT_Tbl):
                table = Table(_element, self.doc)
                t_element = self.parse_table(table)
                self.elements.append(t_element)

            self.position += 1

        # 获取外部的图片 # TODO 这里的图片怎么放到合适的位置呢
        for rid, rel in self.doc.part.rels.items():
            if rel.reltype.endswith("image"):
                if rel.is_external:
                    ref_path = Path(rel.target_ref[8:])
                    ref_img_name = ref_path.name
                    ref_img_ext = ref_path.suffix.lstrip(".")
                    try:
                        with open(ref_path, "rb") as f:
                            _byte = f.read()
                        ref_img_str = base64.b64encode(_byte).decode()
                        img_info = {"img_name": ref_img_name, "img_ext": ref_img_ext, "img_str": ref_img_str}
                        self.doc_images.append(img_info)
                        self.elements.append({"type": "external_image", "image": img_info})
                    except:
                        print(ref_path, "不存在")
                        continue

    def parse_table(self, table):
        """解析表格"""
        # 整个table
        rows_data = []
        for row in table.rows:
            # table 单行
            row_cells = []
            for cell in row.cells:
                # 解析单个cell段落信息
                cell_p_elements = []
                for paragraph in cell.paragraphs:
                    cell_p_element = self.parse_paragraph(paragraph)
                    cell_p_elements.append(cell_p_element)
                row_cells.append(cell_p_elements)
            rows_data.append(row_cells)

        element = {
            "type": "table",
            "data": rows_data
        }
        return element

    def parse_paragraph(self, paragraph: Paragraph):
        """解析段落，包含文本、格式、超链接、图片、文本框"""
        # 解析文本和图片
        p_run_text_parts = []
        p_run_images = []
        for run in paragraph.runs:
            run_data = self.parse_run(run)
            if run_data:
                if run_data["type"] == "text":
                    p_run_text_parts.append(run_data)
                elif run_data["type"] == "image":
                    # run_data = {"type": "image", "images": List[Dict[str, Any]]}
                    p_run_images.extend(run_data["images"])

        # 解析超链接
        hyperlinks = self.get_hyperlinks(paragraph)

        # 解析文本框
        textboxes = self.get_textboxes(paragraph)

        element = {
            "type": "paragraph",
            "style": paragraph.style.name if paragraph.style else "Normal",
            # "alignment": str(paragraph.alignment) if paragraph.alignment else None,
            "text": paragraph.text,
            "runs": p_run_text_parts,
            "hyperlinks": hyperlinks,
            "images": p_run_images,
            "textboxes": textboxes
        }
        return element

    def get_hyperlinks(self, paragraph: Paragraph):
        """获取段落中的超链接"""
        hyperlinks = []
        for child in paragraph._element:
            if child.tag.endswith("hyperlink"):
                # 文本
                link_text = "".join([node.text for node in child.iter() if node.text])
                # 地址
                rel_id = child.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
                if rel_id:
                    try:
                        link_url = paragraph.part.rels[rel_id].target_ref
                        hyperlinks.append({"text": link_text, "url": link_url})
                    except:
                        pass
        return hyperlinks

    def get_textboxes(self, paragraph: Paragraph):
        """提取文本框内容"""
        textboxes = []
        dedup_content = set()
        for child in paragraph._element.iter():
            tag = child.tag
            if "txbxContent" in tag or "textbox" in tag.lower():
                node_texts = []
                for node in child.iter():
                    if node.text and node.text not in node_texts:
                        node_texts.append(node.text)
                text = "".join(node_texts).strip()
                if text not in dedup_content:
                    textboxes.append({"text": text})
                dedup_content.add(text)
        return textboxes

    def parse_run(self, run):
        """解析文本片段（run）格式"""
        # if "graphicData" in run._element.xml or "graphic" in run._element.xml:
        curr_run_images = self.extract_images_from_run(run)
        if curr_run_images:
            return {"type": "image", "images": curr_run_images}

        if not run.text.strip():
            return None

        # 字体格式
        font_format = {
            "bold": run.bold,
            "italic": run.italic,
            "underline": run.underline,
            "font_size": run.font.size.pt if run.font.size else None,
            "font_name": run.font.name,
            "color": f"#{run.font.color.rgb}" if run.font.color and run.font.color.rgb else None,
            "highlight": str(run.font.highlight_color) if run.font.highlight_color else None
        }
        return {"type": "text", "text": run.text, "font_format": font_format}

    def extract_images_from_run(self, run):
        """
        查找run中的图片引用，
        图片在XML中的结构：drawing -> inline/anchor -> graphic -> graphicData -> pic -> blipFill -> blip
        """
        try:
            # 遍历 run 元素 查找包含 blip 的元素
            curr_run_images = []
            for element in run._element.iter():
                tag = element.tag
                if "blip" in tag.lower():
                    embed_key = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
                    embed_id = element.get(embed_key)
                    if embed_id and embed_id in run.part.related_parts:
                        # 根据关系id获取图片
                        image_part = run.part.related_parts[embed_id]
                        # 不考虑文件的扩展名，后期需要将wmf/emf转换为png格式，但是这种格式仅限win打开，linux和macos无法直接打开
                        img_ext = image_part.content_type.split('/')[-1]
                        img_blob = image_part.blob
                        img_str = base64.b64encode(img_blob).decode()
                        self.image_counter += 1
                        img_name = f"image_{self.image_counter}.{img_ext}"
                        img_info = {"img_name": img_name, "img_ext": img_ext, "img_str": img_str}
                        curr_run_images.append(img_info)
                        self.doc_images.append(img_info)

                elif tag == "{urn:schemas-microsoft-com:vml}imagedata":
                    embed_id = element.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
                    if embed_id and embed_id in run.part.related_parts:
                        image_part = run.part.related_parts[embed_id]
                        # 不考虑文件的扩展名，后期需要将wmf/emf转换为png格式，但是这种格式仅限win打开，linux和macos无法直接打开
                        img_ext = image_part.content_type.split('/')[-1]
                        img_blob = image_part.blob
                        img_str = base64.b64encode(img_blob).decode()
                        self.image_counter += 1
                        img_name = f"image_{self.image_counter}.{img_ext}"
                        img_info = {"img_name": img_name, "img_ext": img_ext, "img_str": img_str}
                        curr_run_images.append(img_info)
                        self.doc_images.append(img_info)

            return curr_run_images
        except Exception as e:
            print(f"[Warning]: 图片提取失败 - {repr(e)}")
        return None

    def convert_to_markdown(self):
        """将解析的内容转换为markdown"""
        markdown_lines = []
        for element in self.elements:
            if element["type"] == "paragraph":
                md_text = self._convert_paragraph_to_markdown(element)
            elif element["type"] == "table":
                md_text = self._convert_table_to_markdown(element)
            elif element["type"] == "external_image":
                md_text = f"![media/image]({element['image']['img_name']})"
            else:
                md_text = ""
            if md_text:
                markdown_lines.append(md_text)
        return "\n".join(markdown_lines)

    def _convert_paragraph_to_markdown(self, element):
        """将段落转为markdown"""
        style = element["style"]
        text = element["text"]
        md_text = ""
        if text.strip():
            # 处理标题样式
            if 'Heading 1' in style or 'heading 1' in style.lower():
                md_text = f"# {text}\n"
            elif 'Heading 2' in style or 'heading 2' in style.lower():
                md_text = f"## {text}\n"
            elif 'Heading 3' in style or 'heading 3' in style.lower():
                md_text = f"### {text}\n"
            elif 'Heading 4' in style or 'heading 4' in style.lower():
                md_text = f"#### {text}\n"
            elif 'Heading 5' in style or 'heading 5' in style.lower():
                md_text = f"##### {text}\n"
            elif 'Heading 6' in style or 'heading 6' in style.lower():
                md_text = f"###### {text}\n"
            else:
                # 普通段落，处理格式
                md_text = self._apply_text_format(element)

            if element["hyperlinks"]:
                for link in element["hyperlinks"]:
                    link_text = link["text"]
                    link_url = link["url"]
                    # 替换文本中的链接文字为 Markdown 链接格式
                    md_text = md_text.replace(link_text, f"[{link_text}]({link_url})")

        if element.get("textboxes"):
            for textbox in element["textboxes"]:
                md_text += f"\n> **[textbox]** {textbox['text']}\n"

        if element.get("images"):
            for image in element["images"]:
                img_name = image["img_name"]
                md_text += f"\n![media/image]({img_name})\n"

        if not md_text.strip():
            return ""
        return md_text + "\n"

    def _convert_cell_to_markdown(self, cell):
        """将cell转换为markdown"""
        md_text = ""
        for cell_p in cell:
            md_text += self._convert_paragraph_to_markdown(cell_p)
        return md_text.strip()

    def _convert_table_to_markdown(self, element):
        """将表格转换为markdown"""
        if not element.get("data"):
            return ""
        lines = []
        data = element["data"]
        if len(data) > 0:
            # 第一行
            header = data[0]
            header_text = " | ".join([self._convert_cell_to_markdown(cell) for cell in header])
            lines.append(f"| {header_text} |")
            # 分隔线
            separator = " | ".join(["---" for _ in header])
            lines.append(f"| {separator} |")
            # 剩余行
            for row in data[1:]:
                row_text = " | ".join([self._convert_cell_to_markdown(cell) for cell in row])
                lines.append(f"| {row_text} |")
        return "\n".join(lines) + "\n"

    def _apply_text_format(self, element):
        """正文"""
        result = []
        for run in element["runs"]:
            if run["type"] != "text":
                continue
            text = run["text"]
            fmt = run["font_format"]
            if fmt.get("bold") and fmt.get("italic"):
                text = f"***{text}***"
            elif fmt.get('bold'):
                text = f"**{text}**"
            elif fmt.get('italic'):
                text = f"*{text}*"
            result.append(text)

        return "".join(result) if result else element["text"]


class WordParser(BaseParser):
    def __init__(self, logger, converter:BaseFormatConverter=None):
        super().__init__(logger)
        self.converter = converter
        if self.converter is None:
            self.converter = LibreofficeFormatConverter()

    @property
    def support_file_formats(self) -> List[str]:
        # return [".doc", ".docx", ".docm", ".dot", ".dotx", ".dotm", ".rtf"]
        # 暂时仅支持一下三种格式，其余格式未验证
        return ["doc", "docx", "docm", "rtf"]

    def parse(self, filepath: str | Path, doc_id=None) -> List[LinkDocument]:
        filepath = Path(filepath).absolute()
        origin_filepath = filepath
        suffix = filepath.suffix.lstrip(".")
        if suffix not in self.support_file_formats:
            raise ValueError(f"Word Parser only supports {self.support_file_formats}, but got {suffix}")

        # 若是doc格式，先将doc转为docx
        if suffix != "docx":
            try:
                docx_file = self.converter.convert(
                    input_filepath=filepath,
                    input_format=suffix,
                    output_format="docx",
                    output_suffix="docx"
                )
                filepath = Path(docx_file).absolute()
            except Exception as e:
                self.logger.info(f"[ConvertError]: `doc`->`docx` 转换失败。`{filepath}` - {repr(e)}")
                return []

        try:
            docx_file = DocxParser(filepath)
            # 获取文件属性
            props = docx_file.get_document_properties()
            # 解析所有元素
            docx_file.parse_all()
            markdown = docx_file.convert_to_markdown()
            images = docx_file.doc_images

            link_document = LinkDocument(
                doc_id=doc_id or str(uuid.uuid4()),
                content=markdown.replace("\u3000", " "),
                images=images,
                metadata={"file_name": origin_filepath.name}
            )
            link_document.metadata.update(props)

            return [link_document]

        except Exception as e:
            self.logger.info(f"[ConvertError]: `docx`->`markdown` 转换失败。`{filepath}` - {repr(e)}")
            return []